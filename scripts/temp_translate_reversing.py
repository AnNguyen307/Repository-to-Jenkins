import random
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz
import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

URL = 'https://www.foo.be/cours/dess-20122013/b/Eldad_Eilam-Reversing__Secrets_of_Reverse_Engineering-Wiley%282005%29.pdf'
OUT = Path('translation_output')
OUT.mkdir(exist_ok=True)

print('Downloading the 620-page source...')
raw = requests.get(URL, timeout=240)
raw.raise_for_status()
pdf = fitz.open(stream=raw.content, filetype='pdf')
if pdf.page_count != 620:
    raise RuntimeError(f'Unexpected source page count: {pdf.page_count}')
pages = [p.get_text('text', sort=True) or '' for p in pdf]
pdf.close()
print('Source extracted.')

qxd = re.compile(r'^\s*\d+_\d+\s+\S+\.qxd.*Page\s+\S+\s*$', re.I)
page_number = re.compile(r'^\s*(?:[ivxlcdm]+|\d+)\s*$', re.I)
running = {'contents','foreword','introduction','acknowledgments','foundations','low-level software','windows fundamentals','reversing tools'}
asm = re.compile(r'^\s*(?:[A-Za-z_.$?@][\w.$?@]*:|(?:mov|push|pop|call|ret|cmp|test|j[a-z]+|add|sub|mul|imul|div|idiv|lea|xor|or|and|not|shl|shr|sar|inc|dec|nop|int|rep|lods|stos|scas|cmps)\b)', re.I)
code_start = re.compile(r'^\s*(?:[{}#]|(?:if|for|while|switch|return|class|struct|typedef|public|private|protected|static|void|int|char|long|short|unsigned|signed)\b)', re.I)

def clean_page(text):
    lines=[]
    for ln in text.replace('\u00ad','').replace('￾','').splitlines():
        s=ln.rstrip()
        if qxd.match(s) or page_number.match(s):
            continue
        if s.strip().lower() in running and len(s.strip()) < 40:
            continue
        lines.append(s)
    return '\n'.join(lines).strip()

def codeish(block):
    lines=[x for x in block.splitlines() if x.strip()]
    if not lines:
        return False
    hits=sum(bool(asm.match(x) or code_start.match(x)) for x in lines)
    symbols=sum(block.count(c) for c in '{}[]();=<>\\')
    return hits/len(lines) >= .4 or symbols/max(1,len(block.split())) > .55

def blocks(page):
    result=[]
    for b in re.split(r'\n\s*\n', clean_page(page)):
        b=b.strip()
        if not b:
            continue
        if codeish(b):
            result.append(('code',b))
            continue
        lines=[x.strip() for x in b.splitlines() if x.strip()]
        if len(lines)>=3 and sum(len(x)<58 for x in lines)/len(lines)>.7:
            result.extend(('prose',x) for x in lines)
        else:
            t=' '.join(lines)
            t=re.sub(r'(?<=\w)-\s+(?=\w)','',t)
            result.append(('prose',re.sub(r'\s+',' ',t).strip()))
    return result

def split_text(text, limit=3300):
    if len(text)<=limit:
        return [text]
    sentences=re.split(r'(?<=[.!?])\s+(?=[A-Z0-9“"(])',text)
    out=[]
    cur=''
    for s in sentences:
        if len(s)>limit:
            if cur:
                out.append(cur)
                cur=''
            out.extend(s[i:i+limit] for i in range(0,len(s),limit))
        elif not cur:
            cur=s
        elif len(cur)+1+len(s)<=limit:
            cur+=' '+s
        else:
            out.append(cur)
            cur=s
    if cur:
        out.append(cur)
    return out

token_re=re.compile(r'(https?://\S+|www\.\S+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}|\b0x[0-9A-Fa-f]+\b|\b[A-Z]{2,}[A-Z0-9_]*\b|\b\w+\.(?:exe|dll|sys|c|cpp|h|cs|java|class|pdf|txt|ini|asm|obj|lib)\b)',re.I)

def protect(text):
    vals=[]
    def sub(m):
        vals.append(m.group(0))
        return f'ZXQKEEP{len(vals)-1:04d}QXZ'
    return token_re.sub(sub,text),vals

def restore(text,vals):
    for i,v in enumerate(vals):
        text=re.sub(rf'ZXQ\s*KEEP\s*{i:04d}\s*QXZ',lambda m:v,text,flags=re.I)
    return text

session=requests.Session()
session.headers['User-Agent']='Mozilla/5.0 Chrome/124 Safari/537.36'
endpoint='https://translate.googleapis.com/translate_a/single'

def translate(text):
    if sum(ch.isalpha() for ch in text)<3:
        return text
    masked,vals=protect(text)
    payload={'client':'gtx','sl':'en','tl':'vi','dt':'t','q':masked}
    last=None
    for attempt in range(10):
        try:
            r=session.post(endpoint,data=payload,timeout=60)
            if r.status_code==429:
                raise RuntimeError('rate limit')
            r.raise_for_status()
            data=r.json()
            out=''.join(x[0] for x in data[0] if x and x[0])
            return restore(out,vals)
        except Exception as e:
            last=e
            time.sleep(min(45,2**attempt+random.random()*3))
    raise RuntimeError(f'Translation failed: {last}')

structured=[]
chunks=[]
for p in pages:
    page=[]
    for kind,text in blocks(p):
        pieces=[text] if kind=='code' else split_text(text)
        page.append((kind,pieces))
        if kind=='prose':
            chunks.extend(pieces)
    structured.append(page)
unique=list(dict.fromkeys(chunks))
print(f'Unique prose chunks: {len(unique)}')

translated={}
with ThreadPoolExecutor(max_workers=3) as pool:
    fs={pool.submit(translate,t):t for t in unique}
    for n,f in enumerate(as_completed(fs),1):
        src=fs[f]
        translated[src]=f.result()
        if n%20==0 or n==len(unique):
            print(f'Translated {n}/{len(unique)}')

replacements=[
    ('reverse engineering','kỹ thuật dịch ngược'),
    ('kỹ thuật đảo ngược','kỹ thuật dịch ngược'),
    ('thiết kế ngược','dịch ngược'),
    ('ngôn ngữ lắp ráp','ngôn ngữ hợp ngữ'),
    ('bộ đăng ký','thanh ghi'),
    ('trình tháo rời','trình tháo mã'),
]

def normalize(s):
    for a,b in replacements:
        s=s.replace(a,b).replace(a.title(),b.title())
    return s

translated_pages=[]
for page in structured:
    out=[]
    for kind,pieces in page:
        if kind=='code':
            out.append((kind,pieces[0]))
        else:
            out.append((kind,normalize(' '.join(translated.get(x,x) for x in pieces))))
    translated_pages.append(out)

md=[
    '# DỊCH NGƯỢC: BÍ MẬT CỦA KỸ THUẬT DỊCH NGƯỢC','',
    '**Tác giả nguyên bản:** Eldad Eilam  ',
    '**Nhà xuất bản:** Wiley Publishing, 2005  ',
    '**Bản dịch tiếng Việt có hỗ trợ dịch máy; giữ nguyên mã lệnh, tên API, địa chỉ và tên tệp khi có thể.**','',
    '> Chỉ sử dụng các kỹ thuật trong phạm vi hợp pháp, trên phần mềm và hệ thống mà bạn có quyền phân tích.',''
]
for i,page in enumerate(translated_pages,1):
    md += [f'## Trang PDF gốc {i}','']
    if not page:
        md += ['*Trang chủ yếu chứa hình hoặc không có văn bản trích xuất.*','']
    for kind,text in page:
        md += (['```text',text,'```',''] if kind=='code' else [text,''])
(OUT/'Reversing_Secrets_of_Reverse_Engineering_VI.md').write_text('\n'.join(md),encoding='utf-8')

doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.7)
sec.bottom_margin=Cm(1.7)
sec.left_margin=Cm(2)
sec.right_margin=Cm(1.8)
normal=doc.styles['Normal']
normal.font.name='Arial'
normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
normal.font.size=Pt(10.3)
for name,size in [('Title',24),('Heading 1',18),('Heading 2',13)]:
    st=doc.styles[name]
    st.font.name='Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
    st.font.size=Pt(size)
code=doc.styles.add_style('Code Block',WD_STYLE_TYPE.PARAGRAPH) if 'Code Block' not in doc.styles else doc.styles['Code Block']
code.font.name='Liberation Mono'
code._element.rPr.rFonts.set(qn('w:eastAsia'),'Liberation Mono')
code.font.size=Pt(8)
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('DỊCH NGƯỢC\nBÍ MẬT CỦA KỸ THUẬT DỊCH NGƯỢC')
r.bold=True
r.font.size=Pt(24)
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Eldad Eilam\nBản dịch tiếng Việt').bold=True
doc.add_paragraph('Bản dịch có hỗ trợ dịch máy. Lệnh, mã, tên API, tên tệp và địa chỉ được giữ nguyên khi có thể. Mỗi phần có nhãn trang PDF gốc để đối chiếu.')
doc.add_paragraph('Lưu ý: chỉ phân tích phần mềm và hệ thống mà bạn sở hữu hoặc được cấp phép kiểm tra.')
doc.add_page_break()
for i,page in enumerate(translated_pages,1):
    doc.add_paragraph(f'Trang PDF gốc {i}',style='Heading 2')
    if not page:
        doc.add_paragraph('Trang chủ yếu chứa hình hoặc không có văn bản trích xuất.')
    for kind,text in page:
        doc.add_paragraph(text,style='Code Block' if kind=='code' else None)
    if i%50==0:
        print(f'DOCX: {i}/620')
footer=sec.footer.paragraphs[0]
footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run()
b=OxmlElement('w:fldChar')
b.set(qn('w:fldCharType'),'begin')
ins=OxmlElement('w:instrText')
ins.set(qn('xml:space'),'preserve')
ins.text='PAGE'
e=OxmlElement('w:fldChar')
e.set(qn('w:fldCharType'),'end')
run._r.extend([b,ins,e])
doc.core_properties.title='Dịch ngược: Bí mật của kỹ thuật dịch ngược - Bản dịch tiếng Việt'
doc.core_properties.author='Eldad Eilam (tác giả nguyên bản)'
doc.save(OUT/'Reversing_Secrets_of_Reverse_Engineering_VI.docx')
print('Artifacts created.')
