import random
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz
import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

URL='https://www.foo.be/cours/dess-20122013/b/Eldad_Eilam-Reversing__Secrets_of_Reverse_Engineering-Wiley%282005%29.pdf'
OUT=Path('translation_output')
OUT.mkdir(exist_ok=True)

print('Downloading source...')
r=requests.get(URL,timeout=240)
r.raise_for_status()
pdf=fitz.open(stream=r.content,filetype='pdf')
if pdf.page_count!=620:
    raise RuntimeError(f'Unexpected page count: {pdf.page_count}')
pages=[p.get_text('text',sort=True) or '' for p in pdf]
pdf.close()

qxd=re.compile(r'^\s*\d+_\d+\s+\S+\.qxd.*Page\s+\S+\s*$',re.I)
pageno=re.compile(r'^\s*(?:[ivxlcdm]+|\d+)\s*$',re.I)
running={'contents','foreword','introduction','acknowledgments','foundations','low-level software','windows fundamentals','reversing tools'}
asm=re.compile(r'^\s*(?:[A-Za-z_.$?@][\w.$?@]*:|(?:mov|push|pop|call|ret|cmp|test|j[a-z]+|add|sub|mul|imul|div|idiv|lea|xor|or|and|not|shl|shr|sar|inc|dec|nop|int|rep|lods|stos|scas|cmps)\b)',re.I)
code_start=re.compile(r'^\s*(?:[{}#]|(?:if|for|while|switch|return|class|struct|typedef|public|private|protected|static|void|int|char|long|short|unsigned|signed)\b)',re.I)

def clean(text):
    out=[]
    for ln in text.replace('\u00ad','').replace('￾','').splitlines():
        s=ln.rstrip()
        if qxd.match(s) or pageno.match(s):
            continue
        if s.strip().lower() in running and len(s.strip())<40:
            continue
        out.append(s)
    text='\n'.join(out).strip()
    text=re.sub(r'(?<=\w)-\n(?=\w)','',text)
    return text

def codeish(block):
    lines=[x for x in block.splitlines() if x.strip()]
    if not lines:
        return False
    hits=sum(bool(asm.match(x) or code_start.match(x)) for x in lines)
    symbols=sum(block.count(c) for c in '{}[]();=<>\\')
    return hits/len(lines)>=.4 or symbols/max(1,len(block.split()))>.55

def split_text(text,limit=3300):
    if len(text)<=limit:
        return [text]
    paras=re.split(r'\n\s*\n',text)
    out=[]
    cur=''
    for para in paras:
        para=para.strip()
        if not para:
            continue
        if len(para)>limit:
            if cur:
                out.append(cur);cur=''
            sentences=re.split(r'(?<=[.!?])\s+(?=[A-Z0-9“"(])',para)
            tmp=''
            for s in sentences:
                if len(s)>limit:
                    if tmp: out.append(tmp);tmp=''
                    out.extend(s[i:i+limit] for i in range(0,len(s),limit))
                elif not tmp: tmp=s
                elif len(tmp)+1+len(s)<=limit: tmp+=' '+s
                else: out.append(tmp);tmp=s
            if tmp: out.append(tmp)
        elif not cur:
            cur=para
        elif len(cur)+2+len(para)<=limit:
            cur+='\n\n'+para
        else:
            out.append(cur);cur=para
    if cur:
        out.append(cur)
    return out

token_re=re.compile(r'(https?://\S+|www\.\S+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}|\b0x[0-9A-Fa-f]+\b|\b[A-Z]{2,}[A-Z0-9_]*\b|\b\w+\.(?:exe|dll|sys|c|cpp|h|cs|java|class|pdf|txt|ini|asm|obj|lib)\b)',re.I)

def protect(text):
    values=[]
    # Protect code-like paragraphs first.
    parts=[]
    for block in re.split(r'(\n\s*\n)',text):
        if block.strip() and codeish(block):
            values.append(block)
            parts.append(f'ZXQKEEP{len(values)-1:04d}QXZ')
        else:
            parts.append(block)
    text=''.join(parts)
    def sub(m):
        values.append(m.group(0))
        return f'ZXQKEEP{len(values)-1:04d}QXZ'
    return token_re.sub(sub,text),values

def restore(text,values):
    for i,value in enumerate(values):
        text=re.sub(rf'ZXQ\s*KEEP\s*{i:04d}\s*QXZ',lambda m:value,text,flags=re.I)
    return text

session=requests.Session()
session.headers['User-Agent']='Mozilla/5.0 Chrome/124 Safari/537.36'
endpoint='https://translate.googleapis.com/translate_a/single'

def translate_one(text):
    if sum(ch.isalpha() for ch in text)<3:
        return text
    masked,values=protect(text)
    data={'client':'gtx','sl':'en','tl':'vi','dt':'t','q':masked}
    last=None
    for attempt in range(9):
        try:
            resp=session.post(endpoint,data=data,timeout=60)
            if resp.status_code==429:
                raise RuntimeError('rate limited')
            resp.raise_for_status()
            obj=resp.json()
            result=''.join(seg[0] for seg in obj[0] if seg and seg[0])
            return restore(result,values)
        except Exception as exc:
            last=exc
            time.sleep(min(35,2**attempt+random.random()*2))
    raise RuntimeError(f'Translation failed: {last}')

page_chunks=[split_text(clean(p)) if clean(p) else [] for p in pages]
all_chunks=list(dict.fromkeys(chunk for group in page_chunks for chunk in group))
print(f'Chunks to translate: {len(all_chunks)}')
translations={}
with ThreadPoolExecutor(max_workers=4) as pool:
    futures={pool.submit(translate_one,text):text for text in all_chunks}
    for n,future in enumerate(as_completed(futures),1):
        source=futures[future]
        translations[source]=future.result()
        if n%20==0 or n==len(all_chunks):
            print(f'Translated {n}/{len(all_chunks)}')

terms=[('kỹ thuật đảo ngược','kỹ thuật dịch ngược'),('thiết kế ngược','dịch ngược'),('ngôn ngữ lắp ráp','ngôn ngữ hợp ngữ'),('bộ đăng ký','thanh ghi'),('trình tháo rời','trình tháo mã')]
def normalize(text):
    for a,b in terms:
        text=text.replace(a,b).replace(a.title(),b.title())
    return text
translated_pages=[normalize('\n\n'.join(translations.get(c,c) for c in chunks)) for chunks in page_chunks]

md=['# DỊCH NGƯỢC: BÍ MẬT CỦA KỸ THUẬT DỊCH NGƯỢC','', '**Tác giả nguyên bản:** Eldad Eilam  ','**Nhà xuất bản:** Wiley Publishing, 2005  ','**Bản dịch tiếng Việt có hỗ trợ dịch máy; mã lệnh, API, địa chỉ và tên tệp được giữ nguyên khi có thể.**','', '> Chỉ sử dụng các kỹ thuật trong phạm vi hợp pháp và trên hệ thống được phép phân tích.','']
for i,text in enumerate(translated_pages,1):
    md.extend([f'## Trang PDF gốc {i}',''])
    md.extend([text if text else '*Trang chủ yếu chứa hình hoặc không có văn bản trích xuất.*',''])
(OUT/'Reversing_Secrets_of_Reverse_Engineering_VI.md').write_text('\n'.join(md),encoding='utf-8')

doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.7);sec.bottom_margin=Cm(1.7);sec.left_margin=Cm(2);sec.right_margin=Cm(1.8)
normal=doc.styles['Normal'];normal.font.name='Arial';normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial');normal.font.size=Pt(10.3)
for name,size in [('Title',24),('Heading 1',18),('Heading 2',13)]:
    style=doc.styles[name];style.font.name='Arial';style._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial');style.font.size=Pt(size)
code=doc.styles.add_style('Code Block',WD_STYLE_TYPE.PARAGRAPH) if 'Code Block' not in doc.styles else doc.styles['Code Block']
code.font.name='Liberation Mono';code._element.rPr.rFonts.set(qn('w:eastAsia'),'Liberation Mono');code.font.size=Pt(8)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=p.add_run('DỊCH NGƯỢC\nBÍ MẬT CỦA KỸ THUẬT DỊCH NGƯỢC');run.bold=True;run.font.size=Pt(24)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run('Eldad Eilam\nBản dịch tiếng Việt').bold=True
doc.add_paragraph('Bản dịch có hỗ trợ dịch máy. Lệnh, mã, tên API, tên tệp và địa chỉ được giữ nguyên khi có thể. Mỗi phần có nhãn trang PDF gốc để đối chiếu.')
doc.add_paragraph('Lưu ý: chỉ phân tích phần mềm và hệ thống mà bạn sở hữu hoặc được cấp phép kiểm tra.')
doc.add_page_break()
for i,text in enumerate(translated_pages,1):
    doc.add_paragraph(f'Trang PDF gốc {i}',style='Heading 2')
    if not text:
        doc.add_paragraph('Trang chủ yếu chứa hình hoặc không có văn bản trích xuất.')
    else:
        for block in re.split(r'\n\s*\n',text):
            block=block.strip()
            if block:
                doc.add_paragraph(block,style='Code Block' if codeish(block) else None)
    if i%50==0:
        print(f'DOCX: {i}/620')
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run();begin=OxmlElement('w:fldChar');begin.set(qn('w:fldCharType'),'begin');instr=OxmlElement('w:instrText');instr.set(qn('xml:space'),'preserve');instr.text='PAGE';end=OxmlElement('w:fldChar');end.set(qn('w:fldCharType'),'end');run._r.extend([begin,instr,end])
doc.core_properties.title='Dịch ngược: Bí mật của kỹ thuật dịch ngược - Bản dịch tiếng Việt'
doc.core_properties.author='Eldad Eilam (tác giả nguyên bản)'
doc.save(OUT/'Reversing_Secrets_of_Reverse_Engineering_VI.docx')
print('Artifacts created.')
